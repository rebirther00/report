"""Build V2 of the performance proposal for Daejeon Arts Center.

- Base: template (performance_proposal_template.docx) structure
- Content: filled from wife's V1 (대전문화재단_공연기획제안서_앙상블톤_V1.docx)
- Unknown items kept as [ 와이프 입력: ... ] placeholders with guidance.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


KOREAN_FONT = "맑은 고딕"


def set_korean_font(run, size=None, bold=None, color=None):
    run.font.name = KOREAN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rFonts.set(qn("w:ascii"), KOREAN_FONT)
    rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, size=10.5, bold=False, align=None, color=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_korean_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean_font(r, size=14, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean_font(r, size=12, bold=True, color=RGBColor(0x2E, 0x5C, 0x94))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, *, size=10.5, indent_cm=0.6):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("• ")
    set_korean_font(r, size=size, bold=True)
    r2 = p.add_run(text)
    set_korean_font(r2, size=size)
    return p


def add_placeholder(doc, text, *, indent_cm=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("⬜ ")
    set_korean_font(r, size=10.5, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    r2 = p.add_run(text)
    set_korean_font(r2, size=10.5, color=RGBColor(0xC0, 0x39, 0x2B))
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("※ " + text)
    set_korean_font(r, size=9.5, color=RGBColor(0x55, 0x55, 0x55))
    return p


def style_table_header(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F3A68")
    tcPr.append(shd)
    for p in cell.paragraphs:
        for r in p.runs:
            set_korean_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_table_cell(cell, *, bold=False, size=10.5, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        for r in p.runs:
            set_korean_font(r, size=size, bold=bold)


def set_cell(cell, text, *, bold=False, align=None, size=10.5, color=None, placeholder=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    if placeholder:
        set_korean_font(r, size=size, color=RGBColor(0xC0, 0x39, 0x2B))
    else:
        set_korean_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, header, rows, col_widths_cm=None):
    n_cols = len(header)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    # header
    for ci, h in enumerate(header):
        set_cell(t.rows[0].cells[ci], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_table_header(t.rows[0].cells[ci])
    # rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            is_ph = isinstance(val, str) and val.startswith("[")
            set_cell(t.rows[ri].cells[ci], val if val else "", placeholder=is_ph)
    # widths
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t


def build(out_path: str):
    doc = Document()
    # default style
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

    # page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ======== COVER ========
    add_para(doc, "공연 기획 제안서", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Performance Planning Proposal", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20,
             color=RGBColor(0x55, 0x55, 0x55))
    add_para(doc, "동물 친구들과 떠나는 마법 같은 음악 여행", size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
             color=RGBColor(0x1F, 0x3A, 0x68))
    add_para(doc, "— 동물의 사육제 (Le Carnaval des Animaux) —", size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24,
             color=RGBColor(0x2E, 0x5C, 0x94))

    # cover meta table
    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Light Grid Accent 1"
    set_cell(meta.cell(0, 0), "수신", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(meta.cell(0, 1), "대전예술의전당 음악감독 [ 성명 ] 귀하")
    set_cell(meta.cell(1, 0), "제출", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(meta.cell(1, 1), "앙상블 톤 (Ensemble Tone)")
    set_cell(meta.cell(2, 0), "제출일자", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(meta.cell(2, 1), "2026년 5월 [ DD ]일")
    for row in meta.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(13.0)

    add_para(doc, "", size=10.5, space_after=12)
    add_note(doc, "본 제안서는 2026년 5월 5일 대전예술의전당 어린이날 가족음악회 「동물의 사육제」 관람 후 "
                  "음악감독님의 제안에 따라 작성된 검토용 공연 기획안입니다. 본 양식은 일반적인 공연 기획 제안서 "
                  "표준 구성을 바탕으로 작성되었으며, 예당 측 지정 양식이 있을 경우 해당 양식으로 재구성이 가능합니다.")

    doc.add_page_break()

    # ======== 1. Executive Summary ========
    add_heading1(doc, "1. 요약 (Executive Summary)")
    add_bullet(doc, "제안 배경: 2026년 5월 5일 대전예술의전당 어린이날 가족음악회 「동물의 사육제」 관람 후 "
                    "음악감독님의 제안에 따른 검토용 제안서.")
    add_bullet(doc, "공연 한 줄 콘셉트: \"4K 실사 미디어아트 · 멘티미터 실시간 퀴즈 · 야광 응원봉이 결합된, "
                    "어린이가 끝까지 집중하는 K-가족 클래식 음악회 「동물의 사육제」.\"")
    add_bullet(doc, "핵심 차별점: ① 4K 고화질 실사 동물 영상(캐릭터 일러스트 아님) ② 멘티미터 기반 실시간 양방향 퀴즈 "
                    "③ 곡 분위기에 맞춘 야광 응원봉 공감각 연출 — 음악감독님이 본 단체에 관심을 가지신 핵심 포인트.")
    add_placeholder(doc, "희망 시기 및 공연장: [ 와이프 입력 — 예: 2027년 5월 어린이날 시즌 / 앙상블홀 또는 아트홀 ]")
    add_bullet(doc, "주요 출연진: 앙상블 톤 8~9인 편성 (피아노 5중주 Vn1·Vn2·Va·Vc·Db + 플루트 + 클라리넷 + 타악기), "
                    "해설자 1인.")
    add_bullet(doc, "예상 관객 / 기대 효과: 미취학~초등 자녀 동반 가족 1회 [ 인원 ]명 규모. "
                    "'K-어린이 클래식' 콘텐츠 브랜드 선점 + '과학문화도시 대전' 정체성 강화 + 문화 소외 계층 "
                    "아동 초청을 통한 지역 예술 복지 실현.")

    # ======== 2. 제안 단체 소개 ========
    add_heading1(doc, "2. 제안 단체 소개")

    add_heading2(doc, "2.1 단체 개요")
    add_bullet(doc, "단체명: 앙상블 톤 (Ensemble Tone)")
    add_placeholder(doc, "대표자 / 음악감독: [ 와이프 입력 ]")
    add_placeholder(doc, "설립연도: [ 와이프 입력 ]")
    add_placeholder(doc, "소재지: [ 와이프 입력 — 대전광역시 ○○구 ]")
    add_placeholder(doc, "연락처: [ 와이프 입력 — 대표 연락처 / 이메일 ]")
    add_placeholder(doc, "공식 채널: [ 와이프 입력 — 홈페이지 / 인스타그램 / 유튜브 ]")

    add_heading2(doc, "2.2 단체 연혁 및 활동 이력")
    add_placeholder(doc, "[ 와이프 입력 — 정기연주회 / 초청공연 / 위촉 초연 / 교육 프로그램 등을 시간순으로 기재 ]")
    add_note(doc, "최소 3~5개 대표 활동(연도·공연명·장소·역할)을 채워주시면 신뢰도를 크게 높일 수 있습니다.")

    add_heading2(doc, "2.3 단원 구성")
    add_para(doc, "본 공연 기준 편성: 피아노 5중주 (Vn1, Vn2, Va, Vc, Db) + 플루트 + 클라리넷 + 타악기 "
                  "(총 8~9인 내외) + 해설자 1인.")
    add_placeholder(doc, "단원별 약력: [ 와이프 입력 — 단원명·전공·주요 경력 1줄씩 ]")
    add_note(doc, "별첨 ①(출연자 상세 프로필)에 1인당 1쪽 분량으로 정리할 수 있도록 사진과 약력을 준비해 주세요.")

    add_heading2(doc, "2.4 수상·언론 보도·기관 협력 실적")
    add_placeholder(doc, "[ 와이프 입력 — 수상 내역 / 언론 보도(매체·일자·제목) / 협력 기관(공공·민간) ]")
    add_bullet(doc, "(직전) 2026.05.05 유성선병원 김인홀 — 어린이날 가족음악회 「동물의 사육제」 파일럿 공연 진행.")

    # ======== 3. 기획 의도 및 배경 ========
    add_heading1(doc, "3. 기획 의도 및 배경")

    add_heading2(doc, "3.1 기획 배경")
    add_bullet(doc, "어린이 맞춤형 클래식 콘텐츠의 부재 — 기존 클래식 공연은 어린이가 장시간 집중하기 어려운 정적 구조이고, "
                    "어린이 전용 공연은 음악적 깊이가 부족한 경우가 많음. 양자 간의 공백을 메우는 콘텐츠가 필요.")
    add_bullet(doc, "검증된 파일럿 — 2026년 5월 5일 유성선병원 김인홀에서 진행된 「동물의 사육제」 파일럿 공연에서 "
                    "'테크놀로지(4K 미디어아트, 실시간 퀴즈) + 참여형 연주' 모델의 성공 가능성을 확인. "
                    "장소적 접근성 한계에도 불구하고 유료 관객 유치에 성공.")
    add_bullet(doc, "대전예술의전당과의 확장 — 5월 5일 예당 가족음악회 「동물의 사육제」 흐름과 자연스럽게 연결되며, "
                    "대전 시민 전체를 대상으로 한 고품격 가족 콘텐츠로의 확장이 본 제안의 핵심.")

    add_heading2(doc, "3.2 기획 의도")
    add_bullet(doc, "어린이의 '첫 클래식'을 평생 가는 경험으로 — 학부모 피드백 \"아이들이 끝까지 집중해서 본 첫 클래식 공연\"을 "
                    "예당 무대에서 도시 단위로 재현.")
    add_bullet(doc, "음악적 진정성 + 양방향 참여의 결합 — 생상스 「동물의 사육제」 전곡을 정통 챔버 편성으로 연주하면서, "
                    "동시에 4K 시각·실시간 퀴즈·응원봉 연출로 아이가 공연의 '주체'가 되는 경험을 설계.")
    add_bullet(doc, "지역 예술 생태계 기여 — 단발성 이벤트가 아닌, 매년 어린이날 시즌에 운영 가능한 시리즈 콘텐츠 모델 제안.")

    add_heading2(doc, "3.3 대전 지역 관객 및 예술의전당 정체성과의 정합성")
    add_bullet(doc, "관객 정합성 — 대전은 과학·교육 도시 특성상 어린 자녀를 둔 30~40대 학부모 비중이 높고, "
                    "체험형·교육형 공연 수요가 강함. 본 공연의 양방향 콘셉트와 직접 부합.")
    add_bullet(doc, "예당 정체성 — 가족음악회 시리즈와 정합. 첨단 기술과 순수 예술을 결합한다는 점에서 대전의 "
                    "'과학문화도시' 브랜딩과 시너지 가능.")
    add_bullet(doc, "지역 확장성 — 대전시교육청·구청·문화재단과의 단체 관람 협력, 문화 소외 계층 아동 초청 연계 등 "
                    "공공성 강화 트랙으로 확장 가능.")

    # ======== 4. 공연 개요 ========
    add_heading1(doc, "4. 공연 개요")
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["공연명(가제)", "동물 친구들과 떠나는 마법 같은 음악 여행 — '동물의 사육제'"],
            ["희망 일시", "[ 와이프 입력 — 1안 / 2안 / 3안 (예: 2027.05.05 어린이날 / 2027.05 첫째 주말 / 2026.12 크리스마스 가족 시즌) ]"],
            ["희망 공연장", "[ 와이프 입력 — 아트홀(1,546석) / 앙상블홀(640석) / 원형극장 중 선택. 영상 연출 고려 시 앙상블홀 권장 ]"],
            ["러닝타임", "60~70분 (인터미션 없음) — 어린이 집중력 고려"],
            ["관람등급", "전체 관람가 (만 5세 이상 권장, 영유아 동반 가능 여부는 별도 협의)"],
            ["예상 객석 규모", "[ 와이프 입력 — 예: 1회 600석 × 1~2회차 ]"],
            ["공연 콘셉트(1줄)", "4K 실사 미디어아트 · 멘티미터 실시간 퀴즈 · 야광 응원봉이 결합된 K-가족 클래식 음악회"],
        ],
        col_widths_cm=[3.5, 13.0],
    )

    # ======== 5. 프로그램 구성 ========
    add_heading1(doc, "5. 프로그램 구성")

    add_heading2(doc, "5.1 곡목 리스트")
    add_note(doc, "아래는 표준 구성안입니다. 와이프 단체의 편성/편곡 진척도에 따라 곡 순서·편성·시간을 조정해 주세요. "
                  "특히 ✱ 표시 곡은 단체에서 선곡 보강이 필요한 부분입니다.")
    program_rows = [
        ["1", "C. Saint-Saëns", "동물의 사육제 - I. 서주와 사자왕의 행진", "전체 합주 + 해설", "약 2'00\""],
        ["2", "C. Saint-Saëns", "동물의 사육제 - II. 암탉과 수탉", "Vn1, Vn2, Cl, Pf(편곡)", "약 1'00\""],
        ["3", "C. Saint-Saëns", "동물의 사육제 - III. 야생 당나귀", "Pf 듀오(편곡: Pf+Perc)", "약 0'40\""],
        ["4", "C. Saint-Saëns", "동물의 사육제 - IV. 거북이 / V. 코끼리 / VI. 캥거루", "Vc · Db · Pf 중심", "약 4'30\""],
        ["5", "C. Saint-Saëns", "동물의 사육제 - VII. 수족관", "Fl, Cl, Pf, 현악", "약 2'30\""],
        ["6", "C. Saint-Saëns", "동물의 사육제 - VIII. 노새 / IX. 숲속의 뻐꾸기 / X. 새장", "Fl, Cl, 현악, Pf", "약 4'00\""],
        ["7", "C. Saint-Saëns", "동물의 사육제 - XI. 피아니스트 / XII. 화석", "전체 합주", "약 3'30\""],
        ["8", "C. Saint-Saëns", "동물의 사육제 - XIII. 백조", "Vc + Pf", "약 3'00\""],
        ["9", "C. Saint-Saëns", "동물의 사육제 - XIV. 피날레", "전체 합주", "약 2'00\""],
        ["10 ✱", "A. Menken (편곡)", "「라이온 킹」 메들리 — Circle of Life / Can You Feel the Love Tonight", "전체 합주 + 해설", "약 5'00\""],
        ["11 ✱", "[ 와이프 선곡 ]", "[ 인기 애니메이션 OST 1곡 — 예: 「모아나」 How Far I'll Go / 「겨울왕국」 Let It Go 등 ]", "전체 합주", "약 4'00\""],
        ["12", "—", "앙코르 (관객 응원봉 합주)", "전체 합주 + 관객 참여", "약 3'00\""],
    ]
    add_table(
        doc,
        ["순서", "작곡가", "작품명", "편성/연주자", "연주시간"],
        program_rows,
        col_widths_cm=[1.3, 2.8, 6.5, 4.2, 1.8],
    )
    add_note(doc, "총 러닝타임 약 35~40분 + 해설·전환·관객 참여 시간 포함 시 60~70분 예상. 편성·편곡은 별첨 악보 참조.")

    add_heading2(doc, "5.2 연출 요소")
    add_bullet(doc, "해설·내레이션: 유 — 어린이·가족 친화형 사회자 1인. 곡 사이 30~60초 해설 + 퀴즈 진행. "
                    "(사회자 인선: [ 와이프 입력 ])")
    add_bullet(doc, "영상·시각: 4K 고화질 실사 동물 영상을 대형 스크린에 프로젝션. 캐릭터 일러스트가 아닌 "
                    "곡의 이미지에 부합하는 실사 영상으로 시각적 호기심 극대화. 영상 라이브러리: [ 자체 제작 / 라이선스 ]")
    add_bullet(doc, "관객 참여: 멘티미터(Mentimeter) 기반 실시간 퀴즈. 입장 시 좌석 QR 안내 → 곡 사이 퀴즈 3~5회. "
                    "직전 파일럿 공연 참여율 100% 달성.")
    add_bullet(doc, "공감각 연출: 야광 응원봉을 입장 시 전원에게 배포. 특정 곡(피날레·앙코르)에서 곡 분위기에 맞춰 흔드는 "
                    "역동적 연출로 클래식의 정적 분위기를 탈피.")
    add_bullet(doc, "앙코르: 관객 응원봉이 함께하는 단체 합주 1곡 — 와이프 선곡 진행 중.")

    # ======== 6. 출연진 ========
    add_heading1(doc, "6. 출연진")
    add_table(
        doc,
        ["구분", "성명", "주요 약력"],
        [
            ["음악감독 / 지휘", "[ 와이프 입력 ]", "[ 와이프 입력 — 약력 2~3줄 ]"],
            ["협연자(솔리스트)", "[ 와이프 입력 ]", "[ 와이프 입력 — 특히 XIII. 백조 첼로 솔로, II. 암탉과 수탉 등 솔로 파트 담당자 ]"],
            ["사회 / 내레이션", "[ 와이프 입력 ]", "[ 와이프 입력 — 어린이·가족 음악회 경험 우대 ]"],
            ["악장(Concertmaster)", "[ 와이프 입력 ]", "[ 와이프 입력 ]"],
            ["단원 편성", "8~9인 내외", "Vn1, Vn2, Va, Vc, Db, Fl, Cl, Perc — 단원 명단 별첨 ① 참조"],
        ],
        col_widths_cm=[3.5, 4.0, 9.0],
    )
    add_note(doc, "상세 프로필(사진·약력 1쪽 분량)은 별첨 ①로 첨부합니다.")

    # ======== 7. 무대 및 기술 요구사항 ========
    add_heading1(doc, "7. 무대 및 기술 요구사항")
    add_table(
        doc,
        ["구분", "요구사항"],
        [
            ["무대 배치", "8~9인 챔버 편성. 무대 중앙 후면에 대형 LED 또는 프로젝션 스크린(최소 4K 출력 대응) 필요. 배치도 별첨 ⑤."],
            ["조명", "곡별 분위기 조명 큐 + 무대 영상과의 색온도 연동 필수. 응원봉 발광 효과를 살리기 위한 무대 암전 운영 가능 여부 확인."],
            ["음향", "해설자 무선 마이크 1ch, 영상 음향 연동 라인, 단원 모니터 스피커. 연주는 어쿠스틱 + 보조 마이크 운영."],
            ["영상", "4K 실사 동물 영상 프로젝션(필수). 영상-음악 큐시트 별도 제공. 자막(곡 설명) 표출 검토."],
            ["관객 참여 인프라", "관객 좌석 Wi-Fi(멘티미터 접속용) 또는 LTE 신호 안정성 / 입장 시 야광 응원봉 배포 동선 / 객석 QR 안내물."],
            ["리허설 요구", "영상·음악 싱크 리허설을 포함한 무대 리허설 1.5일 권장. (드레스 리허설 + 테크 리허설 별도 1회)"],
            ["기타", "단원 분장실 / 악기 보관실 / 콘트라베이스·타악기 장비 운반 동선 확보."],
        ],
        col_widths_cm=[3.5, 13.0],
    )

    # ======== 8. 홍보 및 마케팅 계획 ========
    add_heading1(doc, "8. 홍보 및 마케팅 계획")

    add_heading2(doc, "8.1 타깃 관객")
    add_bullet(doc, "주 타깃: 만 5세~초등 저학년 자녀를 동반한 30~40대 학부모 가족 (대전 시내·세종 포함).")
    add_bullet(doc, "부 타깃: 클래식 입문 성인 / 조부모 동반 3대 가족 관객 / 미디어아트·뉴미디어 공연 관심층.")

    add_heading2(doc, "8.2 홍보 채널 및 일정")
    add_bullet(doc, "보도자료: 공연 D-8주 1차(라인업 공개) / D-3주 2차(현장 메이킹·영상 티저) — 대전 지역 매체 + 클래식 전문 매체.")
    add_bullet(doc, "SNS: 단체 인스타그램·유튜브에 4K 실사 동물 영상 티저 시리즈(곡별 15초) 8~10편 운영. "
                    "예당 공식 채널 동시 송출 협의.")
    add_bullet(doc, "협력 매체·플랫폼: 대전시교육청 / 대전·세종 학부모 커뮤니티 / 키즈 체험 플랫폼.")
    add_placeholder(doc, "포스터·리플렛 디자인: [ 자체 / 외주 — 와이프 입력 ]")

    add_heading2(doc, "8.3 예상 관객 동원 계획")
    add_placeholder(doc, "[ 와이프 입력 — 예상 유료 객석 점유율(목표 %), 단체 관람 협력처(유치원·초등학교·문화복지기관), 사전 예매 채널 ]")
    add_bullet(doc, "공공성 트랙: 문화 소외 계층 아동 초청석 ○○석 배정 제안 — 예당·재단 협의 시 별도 운영 가능.")

    # ======== 9. 예산 개요 ========
    add_heading1(doc, "9. 예산 개요")
    add_table(
        doc,
        ["항목", "세부 내역", "금액(원)"],
        [
            ["출연료(지휘·솔리스트·단원)", "[ 와이프 입력 ]", "[ ]"],
            ["편곡·악보 사용료", "라이온 킹·애니메이션 OST 편곡 / 저작권 정산", "[ ]"],
            ["무대 제작·소품", "야광 응원봉 ○○○개 / 영상 디스플레이 운영 / 안내물", "[ ]"],
            ["조명·음향·영상 운영", "4K 영상 제작 또는 라이선스 / 영상-음악 큐시트 운영 인력", "[ ]"],
            ["홍보·디자인·인쇄", "포스터·리플렛·SNS 영상 제작", "[ ]"],
            ["예술감독·기획·운영비", "기획·진행·코디네이션", "[ ]"],
            ["기타(보험·예비비 등)", "공연 책임보험 / 예비비 10%", "[ ]"],
            ["합계", "", "[ ]"],
        ],
        col_widths_cm=[5.0, 8.0, 3.5],
    )

    add_heading2(doc, "9.1 수익 구조 / 분담 방식 (제안 트랙별)")
    add_bullet(doc, "기획공연(예당 주관): 예당이 출연료·제작비를 지급하고 티켓 수익은 예당 귀속. 본 단체는 출연·콘텐츠 제공.")
    add_bullet(doc, "공동기획: 분담 비율 협의(예: 예당 60% / 본 단체 40%) 및 티켓 수익 동일 비율 배분.")
    add_bullet(doc, "대관: 본 단체가 대관료·제작비를 부담하고 티켓 수익을 자체 정산. (예당 측 가이드라인 준수)")
    add_note(doc, "음악감독실 또는 공연기획팀에 사전에 ① 지정 양식 유무 ② 제안 트랙(기획공연 / 공동기획 / 대관)을 확인한 뒤 "
                  "본 섹션의 비중을 조정합니다.")

    # ======== 10. 추진 일정 ========
    add_heading1(doc, "10. 추진 일정")
    add_table(
        doc,
        ["마일스톤", "시점(예시)", "주요 활동"],
        [
            ["제안 검토 / 1차 미팅", "D-6개월", "본 제안서 검토, 음악감독실·기획팀 미팅"],
            ["제안 확정·계약", "D-5개월", "트랙(기획공연/공동기획/대관) 확정, 계약 체결"],
            ["캐스팅 확정", "D-4개월", "지휘·솔리스트·사회자 확정, 단원 스케줄 록"],
            ["편곡·악보·영상 제작", "D-3~D-1개월", "애니메이션 OST 편곡, 4K 실사 영상 라이브러리 구축, 큐시트 작성"],
            ["홍보 시작", "D-8주", "보도자료 1차 배포, 티켓 오픈, SNS 티저 시리즈 시작"],
            ["테크·드레스 리허설", "D-3일", "무대 리허설 1.5일 (영상·음악 싱크 포함)"],
            ["본 공연", "D-Day", "본 공연 + 별도 사진·영상 기록"],
            ["정산·결과 보고", "D+2주", "정산, 결과 보고서, 관객 설문 분석"],
        ],
        col_widths_cm=[4.0, 3.0, 9.5],
    )

    # ======== 11. 기대 효과 ========
    add_heading1(doc, "11. 기대 효과")
    add_bullet(doc, "관객 측면: 어린이의 '첫 클래식'을 평생 가는 경험으로. 양방향 참여로 집중력 유지 → 가족 단위 재방문 동력.")
    add_bullet(doc, "대전예술의전당 측면: 가족 관객 충성도 강화, 차별화된 시즌 콘텐츠 확보, 미디어아트 융합 사례로 기관 브랜드 제고.")
    add_bullet(doc, "제안 단체 측면: 대전 지역 대표 가족음악회 단체로의 포지셔닝, 향후 시리즈화·전국 투어 기반 마련.")
    add_bullet(doc, "지역 문화예술 측면: '과학문화도시 대전' 정체성 강화 / 'K-어린이 클래식' 콘텐츠 브랜드 선점 / "
                    "문화 소외 계층 아동 초청 사업과의 연계로 지역 예술 복지 실현.")

    # ======== 12. 첨부자료 ========
    add_heading1(doc, "12. 첨부자료 (체크리스트)")
    add_bullet(doc, "① 출연자 상세 프로필 (1인 1쪽) — [ 별첨 예정 ]")
    add_bullet(doc, "② 단체 포트폴리오(과거 공연 사진·기록) — [ 별첨 예정 ]")
    add_bullet(doc, "③ 과거 공연 영상 링크 — 2026.05.05 유성선병원 김인홀 파일럿 공연 [ 유튜브 링크 ]")
    add_bullet(doc, "④ 추천서·언론 보도 스크랩 — [ 별첨 예정 ]")
    add_bullet(doc, "⑤ 무대 배치도·기술 라이더 — [ 별첨 예정 ]")
    add_bullet(doc, "⑥ 사업자등록증·단체 등록증 사본 — [ 별첨 예정 ]")

    # ======== 작성 안내 ========
    doc.add_page_break()
    add_heading1(doc, "[ 작성 안내 — 제출 전 삭제 ]")
    add_bullet(doc, "본 양식은 대전예술의전당 공식 양식이 아닌, 일반적인 공연 기획 제안서 표준 구성을 바탕으로 작성된 "
                    "권장 템플릿입니다. 음악감독실 또는 공연기획팀에 사전에 ① 지정 양식 유무 ② 제안 트랙(기획공연 / "
                    "공동기획 / 대관)을 확인한 뒤 본문 비중을 조정해 주세요.")
    add_bullet(doc, "본문 분량은 8~15쪽, 첨부 포함 20~25쪽이 적정합니다.")
    add_bullet(doc, "표지의 음악감독 성명, 단체 정보, 제출 일자 등 [ ] 및 ⬜ 표기 부분을 모두 채워주시기 바랍니다.")
    add_bullet(doc, "⬜(빨강) 표시 = 와이프 직접 입력 필수 항목. 가능하면 단체 내부 자료를 정리해 채워주세요.")

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "V2.docx"
    build(out)
