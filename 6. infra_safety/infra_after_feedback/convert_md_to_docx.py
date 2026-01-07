# -*- coding: utf-8 -*-
"""
마크다운 파일을 DOCX로 변환하는 스크립트
"""
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME = '맑은 고딕'

def set_run_font(run, font_name=FONT_NAME, size=11, bold=False):
    """Run에 한글 폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def parse_markdown_line(line):
    """마크다운 라인 파싱"""
    # 헤딩
    if line.startswith('# '):
        return ('h1', line[2:].strip())
    elif line.startswith('## '):
        return ('h2', line[3:].strip())
    elif line.startswith('### '):
        return ('h3', line[4:].strip())
    elif line.startswith('#### '):
        return ('h4', line[5:].strip())
    
    # 리스트
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        return ('li', line.strip()[2:])
    elif re.match(r'^\d+\.', line.strip()):
        return ('ol', re.sub(r'^\d+\.\s*', '', line.strip()))
    
    # 코드 블록
    elif line.strip().startswith('```'):
        return ('code_block', line.strip()[3:])
    
    # 인용
    elif line.startswith('> '):
        return ('quote', line[2:].strip())
    
    # 테이블 구분자
    elif re.match(r'^\|[-:\s|]+\|$', line.strip()):
        return ('table_sep', None)
    
    # 테이블 행
    elif line.strip().startswith('|') and line.strip().endswith('|'):
        cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
        return ('table_row', cells)
    
    # 일반 텍스트
    elif line.strip():
        return ('p', line.strip())
    
    # 빈 줄
    else:
        return ('empty', None)

def convert_md_to_docx(md_file, docx_file):
    """마크다운을 DOCX로 변환"""
    try:
        # 마크다운 파일 읽기
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # DOCX 문서 생성
        doc = Document()
        
        # 기본 스타일 설정
        style = doc.styles['Normal']
        style.font.name = FONT_NAME
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        
        in_code_block = False
        code_lines = []
        table_rows = []
        in_table = False
        
        for line in lines:
            line_type, content = parse_markdown_line(line)
            
            # 코드 블록 처리
            if line_type == 'code_block':
                if in_code_block:
                    # 코드 블록 종료
                    if code_lines:
                        p = doc.add_paragraph()
                        run = p.add_run('\n'.join(code_lines))
                        set_run_font(run, 'Consolas', 9)
                        code_lines = []
                    in_code_block = False
                else:
                    # 코드 블록 시작
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line.rstrip())
                continue
            
            # 테이블 처리
            if line_type == 'table_row':
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(content)
                continue
            elif in_table and line_type != 'table_sep':
                # 테이블 종료
                if table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(table_rows):
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_data):
                            if j < len(row.cells):
                                cell = row.cells[j]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                run = p.add_run(cell_text)
                                is_bold = i == 0
                                set_run_font(run, size=10, bold=is_bold)
                    
                    table_rows = []
                in_table = False
            
            if line_type == 'table_sep':
                continue
            
            # 헤딩
            if line_type == 'h1':
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, size=16, bold=True)
            elif line_type == 'h2':
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, size=14, bold=True)
            elif line_type == 'h3':
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, size=12, bold=True)
            elif line_type == 'h4':
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, size=11, bold=True)
            
            # 리스트
            elif line_type == 'li' or line_type == 'ol':
                p = doc.add_paragraph(style='List Bullet' if line_type == 'li' else 'List Number')
                run = p.add_run(content)
                set_run_font(run, size=10)
            
            # 인용
            elif line_type == 'quote':
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, size=10)
                p.paragraph_format.left_indent = Inches(0.5)
            
            # 일반 텍스트
            elif line_type == 'p':
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, size=10)
        
        # 마지막 테이블 처리
        if table_rows:
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(table_rows):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        is_bold = i == 0
                        set_run_font(run, size=10, bold=is_bold)
        
        # DOCX 파일 저장
        doc.save(docx_file)
        print(f"[OK] 변환 완료: {os.path.basename(md_file)} -> {os.path.basename(docx_file)}")
        return True
        
    except Exception as e:
        print(f"[ERROR] 변환 실패: {os.path.basename(md_file)}")
        print(f"   오류: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    # 워크스페이스 루트 경로
    workspace_root = r"C:\Users\sharp\OneDrive\문서\gitLocal\report"
    target_dir = os.path.join(workspace_root, "6. infra_safety", "infra_after_feedback")
    
    print(f"대상 디렉토리: {target_dir}")
    print(f"디렉토리 존재 여부: {os.path.exists(target_dir)}")
    
    # 변환할 파일 목록
    files = [
        "개선안.md",
        "상세설명서.md",
        "세부예산.md",
        "도구.md"
    ]
    
    print("=" * 60)
    print("마크다운 → DOCX 변환 시작")
    print("=" * 60)
    
    success_count = 0
    for md_file in files:
        md_path = os.path.join(target_dir, md_file)
        docx_file = md_file.replace('.md', '.docx')
        docx_path = os.path.join(target_dir, docx_file)
        
        print(f"\n[*] 변환 중: {md_file}")
        print(f"   입력: {md_path}")
        print(f"   출력: {docx_path}")
        print(f"   파일 존재: {os.path.exists(md_path)}")
        
        if not os.path.exists(md_path):
            print(f"[!] 파일 없음: {md_file}")
            continue
        
        if convert_md_to_docx(md_path, docx_path):
            success_count += 1
    
    print("\n" + "=" * 60)
    print(f"[OK] 변환 완료: {success_count}/{len(files)}개 파일")
    print(f"   출력 디렉토리: {target_dir}")
    print("=" * 60)
